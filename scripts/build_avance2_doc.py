from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy import text


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from etl.conexion import get_engine


OUTPUT = ROOT / "docs" / "AVANCE 02 IMPLEMENTACION ETL Y DATAMART.docx"
EVIDENCE = ROOT / "docs" / "evidencias_avance2"

NAVY = "10233F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
CYAN = "23A6D5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
GREEN = "178A63"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_width = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for index, width in enumerate(widths_inches):
            cell = row.cells[index]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])
    set_run_font(run, size=9, color=MID_GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 8),
        ("Subtitle", 13, MID_GRAY, 0, 12),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.1

    for current_section in doc.sections:
        header = current_section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = "INTELIGENCIA DE NEGOCIOS | AVANCE 2"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in header_paragraph.runs:
            set_run_font(run, size=8.5, color=MID_GRAY, bold=True)
        footer = current_section.footer
        footer_paragraph = footer.paragraphs[0]
        add_page_number(footer_paragraph)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("INTELIGENCIA DE NEGOCIOS")
    set_run_font(run, size=11, color=CYAN, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Implementacion del ETL y\nconstruccion del DataMart"
    )
    set_run_font(run, size=27, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Caso: Colocaciones de Creditos Mivivienda 2024"
    )
    set_run_font(run, size=15, color=BLUE, bold=True)
    subtitle.paragraph_format.space_after = Pt(42)

    metadata = [
        ("Avance", "Avance 2"),
        ("Integrantes", "Daniel Alonso Correa Chanta\nJose Carlo Castro Franco"),
        ("Universidad", "Universidad Tecnologica del Peru"),
        ("Curso", "Inteligencia de Negocios"),
        ("Docente", "Oscar Eduardo Balcazar Chumacero"),
        ("Lugar y fecha", "Piura, Peru - 2026"),
    ]
    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [1.65, 4.65])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, (label, value) in enumerate(metadata):
        left, right = table.rows[index].cells
        set_cell_shading(left, LIGHT_BLUE)
        left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        left_run = left_p.add_run(label)
        set_run_font(left_run, size=10.5, color=NAVY, bold=True)
        right_p = right.paragraphs[0]
        right_p.paragraph_format.space_after = Pt(0)
        for line_index, line in enumerate(value.split("\n")):
            if line_index:
                right_p.add_run("\n")
            right_run = right_p.add_run(line)
            set_run_font(right_run, size=10.5, color="222222")
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run_font(first, bold=True)
        remaining = paragraph.add_run(text[len(bold_lead):])
        set_run_font(remaining)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        run = paragraph.add_run(item)
        set_run_font(run)


def add_callout(doc, title, text, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, size=10.5, color=color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color="333333")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2:
            for cell in cells:
                set_cell_shading(cell, "F8FAFC")
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index > 0 and len(str(value)) < 25
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = p.add_run(str(value))
            set_run_font(run, size=9.3)
    return table


def add_figure(doc, filename, caption, width=6.45):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(EVIDENCE / filename), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_before = Pt(2)
    caption_paragraph.paragraph_format.space_after = Pt(8)
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=9, color=MID_GRAY, italic=True)


def add_code_block(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "0B1322")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.splitlines()):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.5, color="E6EDF7")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def get_results():
    engine = get_engine()
    with engine.connect() as connection:
        kpis = connection.execute(
            text(
                """
                SELECT
                    SUM(cantidad_creditos) cantidad,
                    ROUND(SUM(monto_credito), 2) monto_total,
                    ROUND(AVG(monto_credito), 2) monto_promedio,
                    ROUND(AVG(tasa_interes), 2) tasa_promedio,
                    MIN(tasa_interes) tasa_minima,
                    MAX(tasa_interes) tasa_maxima,
                    MIN(plazo_meses) plazo_minimo,
                    MAX(plazo_meses) plazo_maximo
                FROM vw_creditos_analitica
                """
            )
        ).mappings().one()
        monthly = connection.execute(
            text(
                """
                SELECT mes_nombre, SUM(cantidad_creditos) cantidad,
                       ROUND(SUM(monto_credito), 2) monto_total
                FROM vw_creditos_analitica
                GROUP BY mes_numero, mes_nombre
                ORDER BY monto_total DESC
                LIMIT 3
                """
            )
        ).mappings().all()
        products = connection.execute(
            text(
                """
                SELECT codigo_producto, SUM(cantidad_creditos) cantidad,
                       ROUND(SUM(monto_credito), 2) monto_total
                FROM vw_creditos_analitica
                GROUP BY codigo_producto
                ORDER BY monto_total DESC
                """
            )
        ).mappings().all()
        departments = connection.execute(
            text(
                """
                SELECT departamento, SUM(cantidad_creditos) cantidad,
                       ROUND(SUM(monto_credito), 2) monto_total
                FROM vw_creditos_analitica
                GROUP BY departamento
                ORDER BY monto_total DESC
                LIMIT 5
                """
            )
        ).mappings().all()
    return kpis, monthly, products, departments


def build_document():
    kpis, monthly, products, departments = get_results()
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Contenido", 1)
    contents = [
        ("1", "Introduccion y objetivo del avance"),
        ("2", "Fuente de datos y perfil de calidad"),
        ("3", "Arquitectura y herramientas utilizadas"),
        ("4", "Implementacion fisica de la base de datos"),
        ("5", "Desarrollo del proceso ETL"),
        ("6", "Carga inicial e incremental"),
        ("7", "Validacion, pruebas y calidad"),
        ("8", "Dashboard web y consumo mediante APIs"),
        ("9", "Resultados de negocio y KPIs"),
        ("10", "Evidencias tecnicas"),
        ("11", "Conclusiones"),
        ("-", "Referencias"),
        ("A", "Anexo: explicacion de archivos Python y comandos"),
    ]
    add_table(doc, ["N.o", "Seccion"], contents, [0.7, 5.8])

    add_heading(doc, "Resumen ejecutivo", 1)
    add_body(
        doc,
        "El presente Avance 2 implementa fisicamente el DataMart definido en "
        "el Avance 1 para analizar las colocaciones de creditos Mivivienda "
        "durante 2024. La solucion utiliza Python para el proceso ETL, MySQL "
        "para staging y modelado dimensional, y Flask con HTML, CSS y "
        "JavaScript para la explotacion visual mediante un dashboard web."
    )
    add_body(
        doc,
        "El archivo origen contiene 13,507 filas. El proceso de calidad "
        "elimino 4,160 filas completamente vacias y 11 duplicados exactos, "
        "obteniendose 9,336 registros validos. La carga inicial inserto los "
        "9,336 hechos y la ejecucion incremental posterior inserto cero "
        "registros, demostrando que la clave tecnica record_hash evita la "
        "duplicacion."
    )
    add_callout(
        doc,
        "Resultado principal",
        "DataMart operativo con cinco dimensiones, una tabla de hechos, "
        "carga incremental, validaciones sin errores y dashboard conectado "
        "directamente a la vista analitica de MySQL.",
        GREEN,
    )

    add_heading(doc, "1. Introduccion y objetivo del avance", 1)
    add_body(
        doc,
        "El Avance 1 definio el problema de negocio, los requerimientos "
        "analiticos, el diccionario inicial y un modelo estrella conceptual. "
        "Este segundo avance transforma esa propuesta en una solucion "
        "funcional, reproducible y verificable."
    )
    add_heading(doc, "1.1 Objetivo", 2)
    add_body(
        doc,
        "Desarrollar la integracion, transformacion y carga de datos para "
        "construir fisicamente un DataMart de colocaciones Mivivienda, "
        "incorporando tablas, relaciones, indices, restricciones, carga "
        "inicial, carga incremental, pruebas de integridad y evidencias de "
        "explotacion de informacion."
    )
    add_heading(doc, "1.2 Alcance implementado", 2)
    add_bullets(
        doc,
        [
            "Base de datos MySQL con staging, dimensiones, hechos y vistas.",
            "Proceso ETL desarrollado en Python y pandas.",
            "Limpieza, conversion de tipos, normalizacion y deduplicacion.",
            "Carga inicial e incremental mediante una clave SHA-256.",
            "Pruebas SQL, pruebas unitarias y pruebas de APIs.",
            "Backend Flask con APIs REST y frontend HTML, CSS y JavaScript.",
            "Dashboard con KPIs, graficos y filtros interactivos.",
        ],
    )

    add_heading(doc, "2. Fuente de datos y perfil de calidad", 1)
    add_body(
        doc,
        "La fuente es el archivo CSV Colocaciones de credito Mivivienda_2024, "
        "publicado por Fondo MIVIVIENDA S.A. El archivo esta separado por "
        "punto y coma y contiene 14 columnas relacionadas con fecha, producto, "
        "ubicacion, institucion financiera, montos, plazo y tasa."
    )
    add_table(
        doc,
        ["Indicador", "Resultado", "Tratamiento ETL"],
        [
            ("Filas fisicas", "13,507", "Lectura completa"),
            ("Filas vacias", "4,160", "Eliminadas"),
            ("Filas no vacias", "9,347", "Evaluadas"),
            ("Duplicados exactos", "11", "Eliminados"),
            ("Filas finales", "9,336", "Cargadas"),
            ("Columnas", "14", "Esquema validado"),
        ],
        [2.3, 1.25, 2.95],
    )
    add_heading(doc, "2.1 Reglas de calidad aplicadas", 2)
    add_bullets(
        doc,
        [
            "Las filas completamente vacias se eliminan antes de convertir tipos.",
            "Las fechas numericas aaaammdd se convierten a DATE.",
            "El ubigeo se conserva como texto CHAR(6).",
            "Los textos se limpian con TRIM, espacios uniformes y mayusculas.",
            "Los montos, tasas y plazos se convierten a valores numericos.",
            "Se rechazan montos, tasas y plazos menores o iguales a cero.",
            "Los duplicados exactos se eliminan antes de la carga.",
        ],
    )

    add_heading(doc, "3. Arquitectura y herramientas utilizadas", 1)
    add_body(
        doc,
        "La solucion se implemento dentro de un solo proyecto. Python ejecuta "
        "el ETL y tambien funciona como backend web mediante Flask. MySQL "
        "almacena el DataMart y el navegador consume APIs JSON para construir "
        "la interfaz con JavaScript y Chart.js."
    )
    add_figure(
        doc,
        "arquitectura.png",
        "Figura 1. Arquitectura implementada para el Avance 2.",
    )
    add_table(
        doc,
        ["Componente", "Tecnologia", "Responsabilidad"],
        [
            ("Fuente", "CSV", "Datos de colocaciones 2024"),
            ("ETL", "Python, pandas", "Extraer, limpiar, transformar y cargar"),
            ("Persistencia", "MySQL", "Staging, DataMart, indices y restricciones"),
            ("Backend", "Flask, SQLAlchemy", "APIs REST y consultas parametrizadas"),
            ("Frontend", "HTML, CSS, JavaScript", "Interfaz, filtros y tabla"),
            ("Graficos", "Chart.js", "Visualizaciones interactivas"),
            ("Pruebas", "unittest y SQL", "Calidad, integridad y APIs"),
        ],
        [1.25, 1.8, 3.45],
    )

    add_heading(doc, "4. Implementacion fisica de la base de datos", 1)
    add_body(
        doc,
        "La base configurada es mivivienda_dm. El script setup_database.py "
        "crea la base y ejecuta los archivos SQL de staging y DataMart. La "
        "tabla de hechos utiliza claves subrogadas hacia cinco dimensiones."
    )
    add_figure(
        doc,
        "modelo_estrella.png",
        "Figura 2. Modelo estrella fisico implementado en MySQL.",
    )
    add_heading(doc, "4.1 Tablas implementadas", 2)
    add_table(
        doc,
        ["Tabla", "Tipo", "Descripcion"],
        [
            ("stg_colocaciones_mivivienda", "Staging", "Area temporal del lote limpio"),
            ("dim_tiempo", "Dimension", "Jerarquia de fecha, anio, trimestre y mes"),
            ("dim_geografia", "Dimension", "Ubigeo, departamento, provincia y distrito"),
            ("dim_producto", "Dimension", "Codigo y descripcion del producto"),
            ("dim_ifi", "Dimension", "Institucion y tipo de IFI"),
            ("dim_plazo", "Dimension", "Meses, rango y categoria de plazo"),
            ("fact_credito", "Hechos", "Medidas de cada credito desembolsado"),
            ("etl_ejecucion", "Auditoria", "Metricas de cada ejecucion ETL"),
        ],
        [2.35, 1.05, 3.1],
    )
    add_heading(doc, "4.2 Relaciones, indices y restricciones", 2)
    add_bullets(
        doc,
        [
            "Cinco claves foraneas relacionan fact_credito con las dimensiones.",
            "Las claves primarias son autoincrementales y subrogadas.",
            "record_hash posee una restriccion UNIQUE para evitar duplicados.",
            "Se crearon indices sobre tiempo, geografia, producto, IFI, plazo y fecha de corte.",
            "Las restricciones CHECK validan montos, tasas, plazos y cantidad_creditos.",
            "Las columnas criticas utilizan NOT NULL para asegurar integridad.",
        ],
    )

    add_heading(doc, "5. Desarrollo del proceso ETL", 1)
    add_body(
        doc,
        "El ETL esta dividido en modulos para separar responsabilidades. "
        "main.py funciona como orquestador y ejecuta Extract, Transform y Load "
        "en el orden correcto."
    )
    add_heading(doc, "5.1 Extract", 2)
    add_body(
        doc,
        "extract.py lee el CSV usando separador punto y coma, conserva "
        "inicialmente todas las columnas como texto y comprueba que el archivo "
        "posea exactamente las 14 columnas esperadas. La opcion "
        "skip_blank_lines=False permite contar las filas vacias como evidencia."
    )
    add_code_block(
        doc,
        """df = pd.read_csv(
    path,
    sep=";",
    dtype="string",
    encoding="utf-8-sig",
    skip_blank_lines=False,
)""",
    )
    add_heading(doc, "5.2 Transform", 2)
    add_body(
        doc,
        "transform.py aplica limpieza, normalizacion, conversion de tipos, "
        "reglas de negocio y deduplicacion. Al finalizar genera record_hash "
        "mediante SHA-256 con los 14 valores normalizados."
    )
    add_callout(
        doc,
        "Clave tecnica",
        "La fuente no incluye un identificador unico de transaccion. "
        "record_hash representa el contenido completo de cada fila y permite "
        "detectar si un credito ya fue cargado.",
    )
    add_heading(doc, "5.3 Load", 2)
    add_body(
        doc,
        "load.py carga primero staging, despues cada dimension y finalmente "
        "fact_credito. Las dimensiones utilizan INSERT IGNORE y claves unicas. "
        "La tabla de hechos se carga mediante JOIN entre staging y las cinco "
        "dimensiones."
    )
    add_code_block(
        doc,
        """staging -> dim_tiempo
        -> dim_geografia
        -> dim_producto
        -> dim_ifi
        -> dim_plazo
        -> fact_credito""",
    )
    add_heading(doc, "5.4 Orquestacion", 2)
    add_code_block(
        doc,
        """df_raw = extract_colocaciones(source_path)
df_clean, metrics = transform(df_raw)
load_staging(df_clean, source_path)
inserted_rows = load_datamart()
record_execution(mode, source_path, metrics, inserted_rows)""",
    )

    add_heading(doc, "6. Carga inicial e incremental", 1)
    add_body(
        doc,
        "La carga inicial reinicia las dimensiones y la tabla de hechos para "
        "construir el DataMart desde cero. La carga incremental conserva los "
        "datos existentes y solo inserta hashes nuevos."
    )
    add_figure(
        doc,
        "ejecucion_etl.png",
        "Figura 3. Evidencia de carga inicial e incremental.",
    )
    add_table(
        doc,
        ["Modo", "Transformadas", "Insertadas", "Resultado"],
        [
            ("Initial", "9,336", "9,336", "Carga completa"),
            ("Incremental", "9,336", "0", "Sin duplicacion"),
        ],
        [1.4, 1.65, 1.65, 1.8],
    )
    add_body(
        doc,
        "El resultado incremental de cero filas no significa que el proceso "
        "haya fallado. Significa que todos los registros del archivo ya "
        "existian en fact_credito y fueron correctamente identificados."
    )

    add_heading(doc, "7. Validacion, pruebas y calidad", 1)
    add_body(
        doc,
        "Las pruebas combinaron consultas SQL de integridad, conciliacion de "
        "conteos, restricciones fisicas y pruebas automatizadas de Python."
    )
    add_figure(
        doc,
        "validaciones.png",
        "Figura 4. Conteos e indicadores de integridad del DataMart.",
    )
    add_heading(doc, "7.1 Pruebas automatizadas", 2)
    add_figure(
        doc,
        "pruebas.png",
        "Figura 5. Resultado de pruebas unitarias y de integracion.",
        width=5.65,
    )
    add_table(
        doc,
        ["Prueba", "Resultado esperado", "Resultado"],
        [
            ("Eliminacion de fila vacia", "Fila descartada", "Correcto"),
            ("Eliminacion de duplicado", "Una sola fila", "Correcto"),
            ("Monto negativo", "Fila invalida", "Correcto"),
            ("Conexion API-MySQL", "HTTP 200", "Correcto"),
            ("Filtro LIMA + NMIV", "Solo datos filtrados", "Correcto"),
            ("Pagina principal", "HTML disponible", "Correcto"),
        ],
        [2.35, 2.25, 1.9],
    )
    add_heading(doc, "7.2 Conciliacion con el origen", 2)
    add_table(
        doc,
        ["Concepto", "Calculo", "Resultado"],
        [
            ("Registros origen", "-", "13,507"),
            ("Menos filas vacias", "13,507 - 4,160", "9,347"),
            ("Menos duplicados", "9,347 - 11", "9,336"),
            ("Registros en hechos", "COUNT(fact_credito)", "9,336"),
            ("Diferencia", "Esperado - cargado", "0"),
        ],
        [2.15, 2.55, 1.8],
    )

    add_heading(doc, "8. Dashboard web y consumo mediante APIs", 1)
    add_body(
        doc,
        "El dashboard principal se implemento con Flask. A diferencia de "
        "Streamlit, esta arquitectura separa explicitamente backend y "
        "frontend, pero ambos se ejecutan dentro del mismo proyecto y con un "
        "solo proceso Python."
    )
    add_heading(doc, "8.1 Backend Flask", 2)
    add_table(
        doc,
        ["Endpoint", "Metodo", "Funcion"],
        [
            ("/", "GET", "Entrega index.html"),
            ("/api/health", "GET", "Verifica conexion con MySQL"),
            ("/api/filtros", "GET", "Lista opciones de filtros"),
            ("/api/dashboard", "GET", "Devuelve KPIs, graficos y detalle"),
        ],
        [2.1, 0.9, 3.5],
    )
    add_heading(doc, "8.2 Frontend", 2)
    add_body(
        doc,
        "index.html define la estructura; styles.css controla el diseno y la "
        "adaptacion responsive; dashboard.js utiliza fetch para consumir las "
        "APIs y actualiza tarjetas, graficos Chart.js y la tabla."
    )
    add_code_block(
        doc,
        """const response = await fetch(`/api/dashboard?${params.toString()}`);
const data = await response.json();
updateKpis(data.kpis);
updateCharts(data);
updateTable(data.detalle);""",
    )
    add_figure(
        doc,
        "dashboard.png",
        "Figura 6. Dashboard conectado al DataMart con filtros interactivos.",
        width=5.75,
    )

    add_heading(doc, "9. Resultados de negocio y KPIs", 1)
    add_table(
        doc,
        ["KPI", "Resultado"],
        [
            ("Cantidad de creditos", f"{int(kpis['cantidad']):,}"),
            ("Monto total colocado", f"S/ {float(kpis['monto_total']):,.2f}"),
            ("Monto promedio", f"S/ {float(kpis['monto_promedio']):,.2f}"),
            ("Tasa promedio", f"{float(kpis['tasa_promedio']):.2f}%"),
            ("Tasa minima / maxima", f"{kpis['tasa_minima']}% / {kpis['tasa_maxima']}%"),
            ("Plazo minimo / maximo", f"{kpis['plazo_minimo']} / {kpis['plazo_maximo']} meses"),
        ],
        [3.25, 3.25],
    )
    add_heading(doc, "9.1 Productos", 2)
    add_table(
        doc,
        ["Producto", "Creditos", "Monto total"],
        [
            (
                row["codigo_producto"],
                f"{int(row['cantidad']):,}",
                f"S/ {float(row['monto_total']):,.2f}",
            )
            for row in products
        ],
        [2.0, 1.5, 3.0],
    )
    add_heading(doc, "9.2 Principales departamentos", 2)
    add_table(
        doc,
        ["Departamento", "Creditos", "Monto total"],
        [
            (
                row["departamento"],
                f"{int(row['cantidad']):,}",
                f"S/ {float(row['monto_total']):,.2f}",
            )
            for row in departments
        ],
        [2.0, 1.5, 3.0],
    )
    add_heading(doc, "9.3 Hallazgos", 2)
    best_month = monthly[0]
    add_bullets(
        doc,
        [
            f"{best_month['mes_nombre']} registro el mayor monto mensual: "
            f"S/ {float(best_month['monto_total']):,.2f}.",
            "NMIV concentro la mayor cantidad y monto de colocaciones.",
            "Lima fue el departamento con mayor concentracion de creditos.",
            "La tasa promedio general fue 10.03% y el plazo maximo fue 300 meses.",
        ],
    )

    add_heading(doc, "10. Evidencias tecnicas", 1)
    add_table(
        doc,
        ["Evidencia solicitada", "Archivo o resultado presentado"],
        [
            ("Scripts SQL", "sql/00 al sql/04"),
            ("Proceso ETL", "etl/extract.py, transform.py, load.py y main.py"),
            ("Tablas cargadas", "9,336 hechos y cinco dimensiones"),
            ("Carga incremental", "Segunda ejecucion: 0 filas nuevas"),
            ("Resultados de consultas", "03_validaciones.sql y 04_consultas_kpi.sql"),
            ("Pruebas", "5 pruebas automatizadas aprobadas"),
            ("Dashboard", "Flask + HTML/CSS/JS en localhost:5000"),
        ],
        [2.4, 4.1],
    )
    add_heading(doc, "10.1 Correspondencia con los entregables", 2)
    add_table(
        doc,
        ["Entregable", "Estado", "Evidencia"],
        [
            ("Base de datos implementada", "Cumplido", "MySQL mivivienda_dm"),
            ("Scripts SQL", "Cumplido", "DDL, vistas, validaciones y KPIs"),
            ("ETL funcional", "Cumplido", "Carga inicial e incremental"),
            ("Hechos y dimensiones", "Cumplido", "Modelo estrella fisico"),
            ("Evidencias de carga", "Cumplido", "Figuras 3 y 4"),
            ("Validaciones realizadas", "Cumplido", "SQL y unittest"),
        ],
        [2.55, 1.25, 2.7],
    )

    add_heading(doc, "11. Conclusiones", 1)
    add_bullets(
        doc,
        [
            "La tabla plana del archivo origen fue transformada en un modelo "
            "estrella con cinco dimensiones y una tabla de hechos.",
            "El proceso ETL resolvio filas vacias, formatos de fecha, "
            "normalizacion de texto, conversion numerica y duplicados.",
            "La clave record_hash permite realizar cargas incrementales sin "
            "repetir registros.",
            "Las pruebas confirmaron que no existen claves foraneas nulas, "
            "montos invalidos ni diferencias entre registros esperados y cargados.",
            "El dashboard web demuestra la explotacion del DataMart mediante "
            "APIs, filtros y visualizaciones interactivas.",
            "La implementacion satisface los entregables tecnicos solicitados "
            "para el Avance 2.",
        ],
    )

    add_heading(doc, "Referencias", 1)
    references = [
        "Fondo MIVIVIENDA S.A. (2025). Colocaciones de los Creditos "
        "Mivivienda 2024. Plataforma Nacional de Datos Abiertos.",
        "Kimball, R., & Ross, M. (2013). The Data Warehouse Toolkit: The "
        "Definitive Guide to Dimensional Modeling (3rd ed.). Wiley.",
        "Pandas Development Team. (2026). pandas documentation.",
        "Pallets Projects. (2026). Flask documentation.",
        "Oracle. (2026). MySQL Reference Manual.",
        "Universidad Tecnologica del Peru. (2026). Guia de Avance 2: "
        "Implementacion del ETL y construccion del DataMart.",
    ]
    add_bullets(doc, references)

    add_heading(doc, "Anexo A. Explicacion de archivos Python", 1)
    file_rows = [
        ("etl/conexion.py", "Lee .env y crea conexiones MySQL con SQLAlchemy."),
        ("etl/setup_database.py", "Crea la base y ejecuta los scripts DDL."),
        ("etl/extract.py", "Lee el CSV y valida sus 14 columnas."),
        ("etl/transform.py", "Limpia, convierte, valida, deduplica y genera hash."),
        ("etl/load.py", "Carga staging, dimensiones, hechos y auditoria."),
        ("etl/main.py", "Orquesta el ETL inicial o incremental."),
        ("web/app.py", "Controlador Flask y endpoints REST."),
        ("web/services/dashboard_service.py", "Consultas y logica del dashboard."),
        ("tests/test_transform.py", "Pruebas de reglas de transformacion."),
        ("tests/test_web_api.py", "Pruebas de pagina, conexion, APIs y filtros."),
    ]
    add_table(doc, ["Archivo", "Responsabilidad"], file_rows, [2.55, 3.95])
    add_heading(doc, "Anexo B. Comandos de ejecucion", 1)
    add_code_block(
        doc,
        """.\\.venv\\Scripts\\python.exe -m etl.setup_database
.\\.venv\\Scripts\\python.exe -m etl.main --mode initial
.\\.venv\\Scripts\\python.exe -m etl.main --mode incremental
.\\.venv\\Scripts\\python.exe -m unittest discover -s tests -v
.\\.venv\\Scripts\\python.exe -m web.app

Dashboard: http://localhost:5000""",
    )

    for section in doc.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    doc.core_properties.title = "Avance 2 - Implementacion ETL y DataMart"
    doc.core_properties.subject = "Inteligencia de Negocios - Mivivienda 2024"
    doc.core_properties.author = "Daniel Alonso Correa Chanta; Jose Carlo Castro Franco"
    doc.core_properties.keywords = "ETL, DataMart, MySQL, Python, Flask, BI"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    build_document()
