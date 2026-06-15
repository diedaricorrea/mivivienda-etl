from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from sqlalchemy import text


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from etl.conexion import get_engine


OUTPUT = ROOT / "docs" / "AVANCE INTEGRADO 01 Y 02 - FORMATO APA 7.docx"
EVIDENCE = ROOT / "docs" / "evidencias_avance2"
FONT = "Times New Roman"


def set_font(run, size=12, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width)
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, settings in edges.items():
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        for key, value in settings.items():
            node.set(qn(f"w:{key}"), str(value))


def add_page_number(header):
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, 12)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    add_page_number(section.header)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(12)
        style.font.color.rgb = None
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.keep_with_next = True

    doc.styles["Title"].font.bold = True
    doc.styles["Title"].font.size = Pt(12)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 3"].font.bold = True
    doc.styles["Heading 3"].font.italic = True
    doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = FONT
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(0)

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(12)
    caption.font.color.rgb = None
    caption.paragraph_format.line_spacing = 2
    caption.paragraph_format.space_after = Pt(0)

    if "APA Reference" not in [style.name for style in doc.styles]:
        reference = doc.styles.add_style("APA Reference", WD_STYLE_TYPE.PARAGRAPH)
        reference.font.name = FONT
        reference._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        reference._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        reference.font.size = Pt(12)
        reference.paragraph_format.line_spacing = 2
        reference.paragraph_format.left_indent = Inches(0.5)
        reference.paragraph_format.first_line_indent = Inches(-0.5)
        reference.paragraph_format.space_after = Pt(0)


def add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Inteligencia de negocios para el sector vivienda: análisis del negocio, "
        "implementación del ETL y construcción del DataMart de colocaciones "
        "Mivivienda 2024"
    )
    set_font(r, 12, bold=True)
    p.paragraph_format.line_spacing = 2

    doc.add_paragraph()
    for text in (
        "Daniel Alonso Correa Chanta",
        "Jose Carlo Castro Franco",
        "Universidad Tecnológica del Perú",
        "Inteligencia de Negocios",
        "Docente: Oscar Eduardo Balcazar Chumacero",
        "Piura, Perú",
        "2026",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 2
    return p


def add_paragraph(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2
        r = p.add_run(item)
        set_font(r)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 2
        r = p.add_run(item)
        set_font(r)


def add_table_label(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Tabla {number}")
    set_font(r, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 2
    p2.paragraph_format.keep_with_next = True
    r2 = p2.add_run(title)
    set_font(r2, italic=True)


def add_apa_table(doc, number, title, headers, rows, widths, note=None):
    add_table_label(doc, number, title)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_font(r, 10, bold=True)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "10", "color": "000000"},
            bottom={"val": "single", "sz": "6", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_font(r, 10)
            borders = {
                "left": {"val": "nil"},
                "right": {"val": "nil"},
                "top": {"val": "nil"},
                "bottom": {"val": "nil"},
            }
            if row_index == len(rows) - 1:
                borders["bottom"] = {"val": "single", "sz": "10", "color": "000000"}
            set_cell_border(cell, **borders)
    if note:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2
        r = p.add_run("Nota. ")
        set_font(r, 10, italic=True)
        r2 = p.add_run(note)
        set_font(r2, 10)
    return table


def add_figure(doc, number, title, filename, width=6.4, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"Figura {number}")
    set_font(r, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 2
    p2.paragraph_format.keep_with_next = True
    r2 = p2.add_run(title)
    set_font(r2, italic=True)
    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.line_spacing = 1
    picture.add_run().add_picture(str(EVIDENCE / filename), width=Inches(width))
    if note:
        p3 = doc.add_paragraph()
        p3.paragraph_format.line_spacing = 2
        r3 = p3.add_run("Nota. ")
        set_font(r3, 10, italic=True)
        r4 = p3.add_run(note)
        set_font(r4, 10)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1
    for index, line in enumerate(code.splitlines()):
        if index:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Courier New"
        r._element.get_or_add_rPr()
        r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
        r._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
        r.font.size = Pt(9)
    for edge in ("top", "bottom", "left", "right"):
        set_cell_border(cell, **{edge: {"val": "single", "sz": "4", "color": "808080"}})


def get_results():
    engine = get_engine()
    with engine.connect() as connection:
        kpis = connection.execute(text("""
            SELECT SUM(cantidad_creditos) cantidad,
                   ROUND(SUM(monto_credito),2) monto_total,
                   ROUND(AVG(monto_credito),2) monto_promedio,
                   ROUND(AVG(tasa_interes),2) tasa_promedio,
                   MIN(tasa_interes) tasa_minima,
                   MAX(tasa_interes) tasa_maxima,
                   MIN(plazo_meses) plazo_minimo,
                   MAX(plazo_meses) plazo_maximo
            FROM vw_creditos_analitica
        """)).mappings().one()
        products = connection.execute(text("""
            SELECT codigo_producto, SUM(cantidad_creditos) cantidad,
                   ROUND(SUM(monto_credito),2) monto_total
            FROM vw_creditos_analitica
            GROUP BY codigo_producto
            ORDER BY monto_total DESC
        """)).mappings().all()
        departments = connection.execute(text("""
            SELECT departamento, SUM(cantidad_creditos) cantidad,
                   ROUND(SUM(monto_credito),2) monto_total
            FROM vw_creditos_analitica
            GROUP BY departamento
            ORDER BY monto_total DESC LIMIT 5
        """)).mappings().all()
    return kpis, products, departments


def build():
    kpis, products, departments = get_results()
    doc = Document()
    configure(doc)
    add_title_page(doc)

    add_heading(doc, "Resumen", 1)
    add_paragraph(
        doc,
        "El presente trabajo desarrolla de manera acumulativa el análisis del "
        "negocio y la implementación física de un DataMart para las colocaciones "
        "de créditos Mivivienda durante 2024. La primera parte define el problema, "
        "los usuarios, los indicadores clave de rendimiento, la fuente de datos y "
        "el modelo dimensional. La segunda parte implementa la base de datos, el "
        "proceso de extracción, transformación y carga, las cargas inicial e "
        "incremental, las pruebas y un dashboard web. La fuente contiene 13,507 "
        "filas; después de eliminar 4,160 filas vacías y 11 duplicados exactos se "
        "cargaron 9,336 hechos. Las pruebas verificaron integridad referencial, "
        "consistencia, calidad, rendimiento y coincidencia con la fuente. La "
        "solución utiliza Python, MySQL y Flask, y permite consultar indicadores "
        "por tiempo, producto, geografía, institución financiera y plazo.",
        indent=False,
    )
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    r = p.add_run(
        "Palabras clave: inteligencia de negocios, ETL, DataMart, modelo estrella, "
        "MySQL, Python."
    )
    set_font(r, italic=True)

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2
    r = p.add_run(
        "Inteligencia de negocios para el sector vivienda: análisis e "
        "implementación del DataMart Mivivienda 2024"
    )
    set_font(r, bold=True)
    add_paragraph(
        doc,
        "El Fondo MIVIVIENDA S.A. promueve el acceso al financiamiento para la "
        "adquisición, construcción y mejoramiento de viviendas. Para este trabajo "
        "se utiliza un conjunto de datos público con las colocaciones realizadas "
        "durante 2024. Aunque la fuente contiene información suficiente para "
        "calcular indicadores, su estructura plana dificulta el análisis repetido "
        "por tiempo, producto, ubicación e institución financiera."
    )
    add_paragraph(
        doc,
        "El modelado dimensional organiza los procesos de negocio alrededor de "
        "hechos medibles y dimensiones descriptivas, facilitando las consultas "
        "analíticas y la comprensión de los datos (Kimball & Ross, 2013). Por ello, "
        "el trabajo integra el análisis realizado en el Avance 1 con la "
        "implementación solicitada en el Avance 2. El documento no presenta dos "
        "informes independientes: conserva las decisiones conceptuales iniciales y "
        "muestra cómo fueron materializadas y verificadas."
    )
    add_paragraph(
        doc,
        "El objetivo general es desarrollar la integración, transformación y carga "
        "de datos para construir físicamente el DataMart de colocaciones "
        "Mivivienda. Los objetivos específicos son implementar tablas, relaciones, "
        "índices y restricciones; desarrollar las etapas Extract, Transform y Load; "
        "realizar cargas inicial e incremental; validar integridad, consistencia, "
        "rendimiento y calidad; y presentar evidencias técnicas y visuales."
    )

    add_heading(doc, "Parte I: Análisis y definición del negocio", 1)
    add_heading(doc, "Definición del problema de negocio", 2)
    add_paragraph(
        doc,
        "La necesidad principal consiste en convertir una tabla plana en una "
        "estructura analítica que permita responder preguntas de gestión de forma "
        "rápida y consistente. El problema no es la ausencia de información, sino "
        "la falta de una organización orientada al análisis. El DataMart debe "
        "permitir conocer cuánto se colocó, cuándo se colocó, mediante qué producto, "
        "en qué ubicación y a través de qué institución financiera."
    )
    add_apa_table(
        doc, 1, "Usuarios finales y necesidades de información",
        ["Usuario", "Necesidad principal"],
        [
            ("Gerencia general", "Revisar KPIs, montos y evolución mensual."),
            ("Analistas de negocio", "Comparar productos, tasas, montos y plazos."),
            ("Área de riesgo", "Evaluar tasas, plazos y concentración."),
            ("Equipo comercial", "Identificar las IFI con mayor colocación."),
            ("Planificación", "Analizar tendencias y concentración geográfica."),
        ],
        [2.0, 4.5],
        "Elaboración propia a partir de los requerimientos del negocio.",
    )

    add_heading(doc, "Objetivos estratégicos e indicadores", 2)
    add_apa_table(
        doc, 2, "Objetivos estratégicos y KPIs",
        ["Objetivo", "Indicadores"],
        [
            ("Analizar colocaciones", "Monto total, cantidad, promedio y evolución mensual."),
            ("Analizar productos", "Cantidad, monto y tasa promedio por producto."),
            ("Analizar IFI", "Ranking por monto y participación por tipo de IFI."),
            ("Analizar geografía", "Créditos por departamento, provincia y distrito."),
            ("Analizar condiciones", "Plazo, tasa y relación cuota inicial/valor."),
        ],
        [2.05, 4.45],
    )

    add_heading(doc, "Proceso de negocio y alcance", 2)
    add_paragraph(
        doc,
        "El proceso principal es el desembolso de un crédito Mivivienda. Cada fila "
        "válida representa un crédito desembolsado por una institución financiera, "
        "en una fecha y ubicación determinadas. En consecuencia, el grano de la "
        "tabla de hechos es un registro por crédito desembolsado."
    )
    add_apa_table(
        doc, 3, "Alcance del DataMart",
        ["Incluye", "No incluye"],
        [
            ("Créditos desembolsados en 2024", "Solicitudes rechazadas o en evaluación"),
            ("Productos NMIV, FCTP y S-CRC", "Datos personales de beneficiarios"),
            ("Análisis geográfico e institucional", "Morosidad, recuperaciones o castigos"),
            ("Montos, tasa, plazo y valor", "Rentabilidad interna de las IFI"),
        ],
        [3.25, 3.25],
    )

    add_heading(doc, "Fuente de datos", 2)
    add_paragraph(
        doc,
        "La fuente corresponde al archivo Colocaciones de los Créditos Mivivienda "
        "2024, publicado por Fondo MIVIVIENDA S.A. (2025). El archivo utiliza punto "
        "y coma como separador y contiene 14 columnas. El CSV funciona como sistema "
        "origen y se carga primero a una tabla staging antes de alimentar el modelo "
        "dimensional."
    )
    add_apa_table(
        doc, 4, "Inventario de la fuente",
        ["Atributo", "Detalle"],
        [
            ("Formato", "CSV separado por punto y coma"),
            ("Cobertura temporal", "3 de enero al 30 de diciembre de 2024"),
            ("Fecha de corte", "8 de enero de 2025"),
            ("Filas físicas", "13,507"),
            ("Filas válidas no vacías", "9,347"),
            ("Columnas", "14"),
        ],
        [2.1, 4.4],
        "Los conteos fueron verificados directamente sobre el archivo entregado.",
    )
    add_apa_table(
        doc, 5, "Diccionario de la tabla origen",
        ["Campo", "Descripción"],
        [
            ("FECHA_DESEMBOLSO", "Fecha del desembolso en formato aaaammdd."),
            ("PRODUCTO", "Código NMIV, FCTP o S-CRC."),
            ("DEPARTAMENTO/PROVINCIA/DISTRITO", "Ubicación de la vivienda."),
            ("UBIGEO", "Código geográfico de seis dígitos."),
            ("IFI/TIPO_IFI", "Institución financiera y clasificación."),
            ("MONTO_CREDITO", "Monto del crédito desembolsado."),
            ("MONTO_CUOTA_INICIAL", "Cuota inicial aportada."),
            ("PLAZOS", "Plazo del crédito en meses."),
            ("TASA", "Tasa de interés anual."),
            ("MONTO_VALOR_VIVIENDA", "Valor total de la vivienda."),
            ("FECHA_CORTE", "Fecha de corte del conjunto de datos."),
        ],
        [2.35, 4.15],
    )

    add_heading(doc, "Problemas de calidad identificados", 2)
    add_apa_table(
        doc, 6, "Problemas de calidad y tratamiento",
        ["Problema", "Tratamiento ETL"],
        [
            ("4,160 filas completamente vacías", "Eliminar antes de convertir tipos."),
            ("Fechas numéricas aaaammdd", "Convertir a DATE."),
            ("UBIGEO susceptible de conversión numérica", "Mantener como CHAR(6)."),
            ("Textos con formatos variables", "Aplicar TRIM, espacios uniformes y mayúsculas."),
            ("Datos repetidos", "Eliminar duplicados exactos y generar hash."),
            ("Redundancia geográfica e institucional", "Separar en dimensiones."),
        ],
        [3.0, 3.5],
    )

    add_heading(doc, "Requerimientos analíticos", 2)
    add_apa_table(
        doc, 7, "Preguntas de negocio",
        ["N.º", "Pregunta"],
        [
            ("1", "¿Cuál fue el monto total desembolsado por mes?"),
            ("2", "¿Qué producto tuvo mayor cantidad de créditos?"),
            ("3", "¿Qué departamentos concentran las colocaciones?"),
            ("4", "¿Cuál es el ranking de IFI por monto colocado?"),
            ("5", "¿Cómo varía la tasa por producto y tipo de IFI?"),
            ("6", "¿Cuántos créditos pertenecen a cada categoría de plazo?"),
            ("7", "¿Qué trimestre registró el mayor monto?"),
        ],
        [0.7, 5.8],
    )

    add_heading(doc, "Modelo conceptual inicial", 2)
    add_paragraph(
        doc,
        "Se seleccionó un modelo estrella porque ofrece una estructura simple para "
        "consultas agregadas y es coherente con una fuente desnormalizada. El modelo "
        "considera una tabla de hechos central y cinco dimensiones: tiempo, "
        "geografía, producto, institución financiera y plazo. Esta decisión sigue "
        "los principios de modelado dimensional propuestos por Kimball y Ross (2013)."
    )
    add_apa_table(
        doc, 8, "Dimensiones del modelo conceptual",
        ["Dimensión", "Atributos principales"],
        [
            ("DIM_TIEMPO", "Fecha, año, semestre, trimestre, mes y día."),
            ("DIM_GEOGRAFIA", "Ubigeo, departamento, provincia y distrito."),
            ("DIM_PRODUCTO", "Código, nombre y descripción."),
            ("DIM_IFI", "Nombre y tipo de institución."),
            ("DIM_PLAZO", "Meses, años, rango y categoría."),
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "Parte II: Implementación del ETL y construcción del DataMart", 1)
    add_heading(doc, "Metodología y herramientas", 2)
    add_paragraph(
        doc,
        "El enunciado presenta PostgreSQL, SQL Server y Oracle como herramientas "
        "posibles, y menciona SSIS, Pentaho y Talend como alternativas ETL. La "
        "expresión herramientas posibles no restringe la implementación a dichas "
        "opciones. Se utilizó MySQL porque el Avance 1 ya lo establecía como destino "
        "del DataMart, y Python con pandas porque permite implementar explícitamente "
        "las etapas de extracción, transformación y carga. MySQL proporciona claves "
        "primarias, foráneas, índices, restricciones y vistas; pandas ofrece lectura, "
        "limpieza y transformación tabular (Oracle, 2026; pandas development team, "
        "2026)."
    )
    add_figure(
        doc, 1, "Arquitectura de la solución implementada",
        "arquitectura.png", 6.4,
        "El navegador consume APIs Flask; el CSV solo es procesado por el ETL.",
    )

    add_heading(doc, "Implementación de la base de datos", 2)
    add_paragraph(
        doc,
        "La base mivivienda_dm se crea mediante setup_database.py, que ejecuta los "
        "scripts 01_staging.sql y 02_datamart.sql. Se implementaron staging, cinco "
        "dimensiones, la tabla fact_credito, la tabla de auditoría etl_ejecucion y "
        "dos vistas analíticas."
    )
    add_figure(
        doc, 2, "Modelo físico del DataMart",
        "modelo_estrella.png", 6.4,
        "Las relaciones son de uno a muchos desde cada dimensión hacia fact_credito.",
    )
    add_apa_table(
        doc, 9, "Objetos físicos implementados",
        ["Objeto", "Función"],
        [
            ("stg_colocaciones_mivivienda", "Recibe el lote transformado."),
            ("dim_tiempo", "Organiza la jerarquía temporal."),
            ("dim_geografia", "Almacena ubicaciones únicas por ubigeo."),
            ("dim_producto", "Describe los productos."),
            ("dim_ifi", "Describe las instituciones financieras."),
            ("dim_plazo", "Clasifica el plazo en meses y categorías."),
            ("fact_credito", "Almacena medidas y claves foráneas."),
            ("etl_ejecucion", "Registra métricas de cada carga."),
            ("vw_creditos_analitica", "Integra hechos y dimensiones para BI."),
        ],
        [2.7, 3.8],
    )
    add_heading(doc, "Relaciones, índices y restricciones", 3)
    add_bullets(doc, [
        "Cinco claves foráneas conectan fact_credito con las dimensiones.",
        "Las claves subrogadas son autoincrementales.",
        "record_hash posee una restricción UNIQUE.",
        "Se crearon índices para fecha, ubigeo, producto y claves foráneas.",
        "Las restricciones CHECK controlan montos, tasas, plazos y cantidad.",
        "Las columnas críticas utilizan NOT NULL.",
    ])

    add_heading(doc, "Desarrollo del proceso ETL", 2)
    add_heading(doc, "Extract", 3)
    add_paragraph(
        doc,
        "extract.py lee el archivo con separador punto y coma, codificación "
        "UTF-8 y todas las columnas como texto. También valida que estén presentes "
        "las 14 columnas esperadas. La lectura conserva las filas vacías para poder "
        "contarlas y documentar su eliminación."
    )
    add_code(doc, """df = pd.read_csv(
    path,
    sep=";",
    dtype="string",
    encoding="utf-8-sig",
    skip_blank_lines=False,
)""")
    add_heading(doc, "Transform", 3)
    add_paragraph(
        doc,
        "transform.py reemplaza cadenas vacías por valores nulos, elimina filas "
        "vacías, normaliza textos, convierte fechas y campos numéricos, valida "
        "reglas de negocio y elimina duplicados. Después genera record_hash con "
        "SHA-256 sobre los 14 valores normalizados. Esta clave técnica compensa la "
        "ausencia de un identificador de transacción en la fuente."
    )
    add_heading(doc, "Load", 3)
    add_paragraph(
        doc,
        "load.py carga staging, dimensiones y hechos en ese orden. Las dimensiones "
        "se insertan antes de fact_credito para disponer de las claves subrogadas. "
        "La integración se realiza mediante JOIN entre staging y las cinco "
        "dimensiones. Finalmente, record_execution registra el modo, los conteos y "
        "el estado de la ejecución."
    )

    add_heading(doc, "Carga inicial e incremental", 2)
    add_paragraph(
        doc,
        "La carga inicial reinicia dimensiones y hechos y construye el DataMart "
        "desde cero. La carga incremental mantiene la información existente y "
        "utiliza record_hash para insertar únicamente filas nuevas."
    )
    add_figure(
        doc, 3, "Ejecución de las cargas inicial e incremental",
        "ejecucion_etl.png", 6.4,
        "La carga inicial insertó 9,336 hechos y la incremental posterior insertó cero.",
    )
    add_apa_table(
        doc, 10, "Resultados de carga",
        ["Modo", "Leídas", "Vacías", "Duplicados", "Transformadas", "Insertadas"],
        [
            ("Inicial", "13,507", "4,160", "11", "9,336", "9,336"),
            ("Incremental", "13,507", "4,160", "11", "9,336", "0"),
        ],
        [1.1, 1.0, 1.0, 1.05, 1.25, 1.1],
    )

    add_heading(doc, "Validación y pruebas", 2)
    add_heading(doc, "Integridad de datos", 3)
    add_paragraph(
        doc,
        "Se comprobó que las 9,336 filas de staging tienen correspondencia en "
        "fact_credito, que todos los hashes son únicos, que no existen claves "
        "foráneas nulas y que no existen hechos huérfanos. Las restricciones de "
        "MySQL complementan estas consultas al impedir valores incompatibles."
    )
    add_figure(
        doc, 4, "Resultados de integridad y calidad",
        "validaciones.png", 6.4,
        "Los conteos fueron obtenidos directamente de MySQL.",
    )
    add_heading(doc, "Consistencia y comparación con el origen", 3)
    add_apa_table(
        doc, 11, "Conciliación entre origen y DataMart",
        ["Concepto", "Cálculo", "Resultado"],
        [
            ("Filas físicas", "-", "13,507"),
            ("Menos filas vacías", "13,507 - 4,160", "9,347"),
            ("Menos duplicados", "9,347 - 11", "9,336"),
            ("Hechos cargados", "COUNT(fact_credito)", "9,336"),
            ("Diferencia", "Esperado - cargado", "0"),
        ],
        [2.2, 2.5, 1.8],
    )
    add_heading(doc, "Calidad de información", 3)
    add_paragraph(
        doc,
        "La calidad se evaluó mediante nulos críticos, montos positivos, fechas "
        "válidas, formato de ubigeo, duplicados y consistencia de claves. El "
        "resultado fue cero claves foráneas nulas, cero montos inválidos y 9,336 "
        "hashes únicos."
    )
    add_heading(doc, "Rendimiento", 3)
    add_paragraph(
        doc,
        "Se ejecutó 20 veces la consulta completa del servicio del dashboard, que "
        "incluye KPIs, serie mensual, productos, departamentos, instituciones y "
        "100 registros de detalle. El promedio fue 186.91 ms, la mediana 181.67 "
        "ms, el mínimo 171.93 ms y el máximo 227.49 ms. Para el volumen académico "
        "de 9,336 hechos, los tiempos son adecuados para una interacción web local."
    )
    add_apa_table(
        doc, 12, "Prueba de rendimiento del dashboard",
        ["Ejecuciones", "Promedio", "Mediana", "Mínimo", "Máximo"],
        [("20", "186.91 ms", "181.67 ms", "171.93 ms", "227.49 ms")],
        [1.3, 1.3, 1.3, 1.3, 1.3],
        "Medición local realizada sobre DashboardService conectado a MySQL.",
    )
    add_heading(doc, "Pruebas automatizadas", 3)
    add_figure(
        doc, 5, "Ejecución de pruebas automatizadas",
        "pruebas.png", 5.8,
        "Las cinco pruebas finalizaron correctamente.",
    )
    add_apa_table(
        doc, 13, "Cobertura de pruebas automatizadas",
        ["Prueba", "Resultado"],
        [
            ("Eliminación de filas vacías", "Correcto"),
            ("Eliminación de duplicados", "Correcto"),
            ("Rechazo de valores de negocio inválidos", "Correcto"),
            ("Conexión de la API con MySQL", "Correcto"),
            ("Aplicación de filtros LIMA y NMIV", "Correcto"),
            ("Entrega de la página HTML", "Correcto"),
        ],
        [4.6, 1.9],
    )

    add_heading(doc, "Verificación de KPIs", 2)
    add_apa_table(
        doc, 14, "Indicadores generales del DataMart",
        ["Indicador", "Resultado"],
        [
            ("Cantidad de créditos", f"{int(kpis['cantidad']):,}"),
            ("Monto total", f"S/ {float(kpis['monto_total']):,.2f}"),
            ("Monto promedio", f"S/ {float(kpis['monto_promedio']):,.2f}"),
            ("Tasa promedio", f"{float(kpis['tasa_promedio']):.2f}%"),
            ("Tasa mínima y máxima", f"{kpis['tasa_minima']}% y {kpis['tasa_maxima']}%"),
            ("Plazo mínimo y máximo", f"{kpis['plazo_minimo']} y {kpis['plazo_maximo']} meses"),
        ],
        [3.3, 3.2],
    )
    add_apa_table(
        doc, 15, "Resultados por producto",
        ["Producto", "Créditos", "Monto total"],
        [
            (row["codigo_producto"], f"{int(row['cantidad']):,}", f"S/ {float(row['monto_total']):,.2f}")
            for row in products
        ],
        [2.0, 1.5, 3.0],
    )
    add_apa_table(
        doc, 16, "Principales departamentos",
        ["Departamento", "Créditos", "Monto total"],
        [
            (row["departamento"], f"{int(row['cantidad']):,}", f"S/ {float(row['monto_total']):,.2f}")
            for row in departments
        ],
        [2.0, 1.5, 3.0],
    )

    add_heading(doc, "Dashboard web", 2)
    add_paragraph(
        doc,
        "La capa de presentación se implementó con Flask, HTML, CSS, JavaScript "
        "y Chart.js. Flask sirve la página y expone las rutas /api/health, "
        "/api/filtros y /api/dashboard. JavaScript utiliza fetch para solicitar "
        "JSON y actualizar tarjetas, gráficos y tabla. Flask es un framework "
        "flexible que permite organizar controladores y servicios sin exigir una "
        "estructura específica (Pallets Projects, 2026)."
    )
    add_figure(
        doc, 6, "Dashboard de colocaciones Mivivienda 2024",
        "dashboard.png", 6.4,
        "La interfaz consulta la vista analítica de MySQL y permite filtrar por departamento, producto y tipo de IFI.",
    )

    add_heading(doc, "Evidencias técnicas y cumplimiento", 2)
    add_apa_table(
        doc, 17, "Matriz de cumplimiento de los requisitos del Avance 2",
        ["Requisito", "Implementación y evidencia", "Estado"],
        [
            ("Creación de tablas", "01_staging.sql y 02_datamart.sql", "Cumplido"),
            ("Relaciones", "Cinco claves foráneas en fact_credito", "Cumplido"),
            ("Índices", "Índices de staging, hechos y vistas", "Cumplido"),
            ("Restricciones", "PK, FK, UNIQUE, NOT NULL y CHECK", "Cumplido"),
            ("Extract", "etl/extract.py", "Cumplido"),
            ("Transform", "etl/transform.py", "Cumplido"),
            ("Load", "etl/load.py y etl/main.py", "Cumplido"),
            ("Limpieza y formatos", "Filas vacías, fechas, textos y números", "Cumplido"),
            ("Duplicados", "drop_duplicates y record_hash", "Cumplido"),
            ("Integración", "Cinco dimensiones y una tabla de hechos", "Cumplido"),
            ("Carga inicial", "9,336 filas insertadas", "Cumplido"),
            ("Carga incremental", "0 duplicados en segunda carga", "Cumplido"),
            ("Integridad", "Cero FK nulas y cero hechos huérfanos", "Cumplido"),
            ("Consistencia", "Diferencia origen/DataMart igual a cero", "Cumplido"),
            ("Rendimiento", "20 consultas; promedio 186.91 ms", "Cumplido"),
            ("Calidad", "Cero montos inválidos y hashes únicos", "Cumplido"),
            ("Verificación de KPIs", "04_consultas_kpi.sql y dashboard", "Cumplido"),
            ("Evidencias", "Figuras 1 a 6, tablas y scripts", "Cumplido"),
        ],
        [2.05, 3.55, 0.9],
        "La matriz relaciona cada actividad solicitada con un artefacto verificable del proyecto.",
    )

    add_heading(doc, "Discusión", 1)
    add_paragraph(
        doc,
        "La implementación confirma que una fuente compuesta por una sola tabla "
        "puede transformarse en un DataMart útil cuando se define correctamente el "
        "grano y se separan los contextos descriptivos. El modelo estrella reduce "
        "la repetición lógica en las consultas y permite que los indicadores se "
        "calculen sobre una estructura estable."
    )
    add_paragraph(
        doc,
        "La ausencia de un identificador único en la fuente constituyó la principal "
        "limitación técnica. El hash del contenido permite idempotencia frente al "
        "mismo archivo; sin embargo, si dos operaciones reales distintas tuvieran "
        "exactamente los mismos 14 valores, serían interpretadas como duplicadas. "
        "En una implementación institucional sería preferible que la fuente "
        "proporcione un identificador transaccional."
    )
    add_paragraph(
        doc,
        "La arquitectura Flask facilita la explicación del proyecto porque separa "
        "las rutas HTTP, el servicio de consultas y el frontend. No obstante, el "
        "objetivo central del avance es el ETL y el DataMart; el dashboard actúa "
        "como evidencia de explotación y no reemplaza las validaciones del modelo."
    )

    add_heading(doc, "Conclusiones", 1)
    add_numbered(doc, [
        "El análisis del Avance 1 se materializó en un modelo estrella físico con cinco dimensiones y una tabla de hechos.",
        "El proceso ETL trató filas vacías, formatos, textos, valores numéricos y duplicados, y cargó 9,336 hechos válidos.",
        "La carga incremental fue idempotente para el archivo evaluado y no insertó registros repetidos.",
        "Las pruebas demostraron integridad referencial, consistencia con la fuente, calidad de información y tiempos de consulta adecuados.",
        "Los KPIs y el dashboard responden las preguntas de negocio definidas inicialmente.",
        "Las evidencias y la matriz de cumplimiento demuestran que se cubrieron las actividades y entregables solicitados para el Avance 2.",
    ])

    add_heading(doc, "Referencias", 1)
    references = [
        "Fondo MIVIVIENDA S.A. (2025). Colocaciones de los créditos Mivivienda 2024 [Conjunto de datos]. Plataforma Nacional de Datos Abiertos.",
        "Kimball, R., & Ross, M. (2013). The data warehouse toolkit: The definitive guide to dimensional modeling (3rd ed.). Wiley.",
        "Oracle. (2026). MySQL reference manual. https://dev.mysql.com/doc/",
        "Pallets Projects. (2026). Flask documentation. https://flask.palletsprojects.com/",
        "pandas development team. (2026). pandas documentation. https://pandas.pydata.org/docs/",
        "Universidad Tecnológica del Perú. (2026). Guía de Avance 2: Implementación del ETL y construcción del DataMart [Material del curso].",
    ]
    for reference in references:
        p = doc.add_paragraph(style="APA Reference")
        r = p.add_run(reference)
        set_font(r)

    add_heading(doc, "Apéndice A", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Archivos principales del proyecto")
    set_font(r, bold=True)
    add_apa_table(
        doc, 18, "Archivos y responsabilidades",
        ["Archivo", "Responsabilidad"],
        [
            ("etl/conexion.py", "Configuración y conexiones MySQL."),
            ("etl/setup_database.py", "Creación automática de la base."),
            ("etl/extract.py", "Lectura y validación del CSV."),
            ("etl/transform.py", "Limpieza, conversión, reglas y hash."),
            ("etl/load.py", "Carga de staging, dimensiones y hechos."),
            ("etl/main.py", "Orquestación inicial e incremental."),
            ("web/app.py", "Rutas Flask y respuestas JSON/HTML."),
            ("web/services/dashboard_service.py", "Consultas analíticas y filtros."),
            ("sql/03_validaciones.sql", "Pruebas de integridad y conciliación."),
            ("sql/04_consultas_kpi.sql", "Consultas de indicadores."),
        ],
        [2.8, 3.7],
    )

    add_heading(doc, "Apéndice B", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Comandos de ejecución")
    set_font(r, bold=True)
    add_code(doc, """.\\.venv\\Scripts\\python.exe -m etl.setup_database
.\\.venv\\Scripts\\python.exe -m etl.main --mode initial
.\\.venv\\Scripts\\python.exe -m etl.main --mode incremental
.\\.venv\\Scripts\\python.exe -m unittest discover -s tests -v
.\\.venv\\Scripts\\python.exe -m web.app

Dashboard: http://localhost:5000""")

    doc.core_properties.title = (
        "Inteligencia de negocios para el sector vivienda: Avances 1 y 2"
    )
    doc.core_properties.author = (
        "Daniel Alonso Correa Chanta; Jose Carlo Castro Franco"
    )
    doc.core_properties.subject = "Documento acumulativo en formato APA 7"
    doc.core_properties.keywords = "BI, ETL, DataMart, APA, MySQL, Python, Flask"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Documento APA generado: {OUTPUT}")


if __name__ == "__main__":
    build()
